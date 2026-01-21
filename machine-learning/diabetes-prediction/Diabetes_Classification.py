import numpy as np
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier, AdaBoostClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.svm import SVC
from sklearn.metrics import classification_report, roc_auc_score, roc_curve, confusion_matrix, ConfusionMatrixDisplay
from sklearn.preprocessing import PolynomialFeatures, StandardScaler
import matplotlib.pyplot as plt
from imblearn.over_sampling import SMOTE
from docx import Document
from docx.shared import Inches
import os

# List of different random seeds for reproducibility
seeds = [42, 123, 456]  # You can add more seeds here

# List of base classifiers to try
classifiers = {
    "RandomForest": RandomForestClassifier(random_state=42),
    "GradientBoosting": GradientBoostingClassifier(random_state=42),
    "AdaBoost": AdaBoostClassifier(random_state=42, algorithm='SAMME'),  
    "LogisticRegression": LogisticRegression(random_state=42, max_iter=200),
    "SVC": SVC(probability=True, random_state=42)
}


# Initialize a document for overall results
overall_doc = Document()
overall_doc.add_heading('Diabetes Classification Results Overview', level=1)

# Define the supervised learning function
def run_supervised_learning_process(classifier_name, classifier, seed, base_folder):
    print(f"\nRunning supervised learning process for {classifier_name} with seed {seed}...\n")

    # Set the random seed for reproducibility
    np.random.seed(seed)

    # Load the dataset
    df = pd.read_csv('diabetes.csv')

    # Extract features and target variable
    X = df.drop(columns=['Outcome'])
    y = df['Outcome']

    # Initialize SMOTE to balance the dataset
    smote = SMOTE(random_state=seed)
    X_resampled, y_resampled = smote.fit_resample(X, y)

    # Split the dataset into training and testing sets
    X_train, X_test, y_train, y_test = train_test_split(X_resampled, y_resampled, test_size=0.3, random_state=seed)

    # Feature Engineering: Create polynomial features
    poly = PolynomialFeatures(degree=2, interaction_only=True, include_bias=False)
    X_train_poly = poly.fit_transform(X_train)
    X_test_poly = poly.transform(X_test)

    # Optional: Scale features
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train_poly)
    X_test_scaled = scaler.transform(X_test_poly)

    # Train the specified classifier on the labeled portion of the training data
    classifier.fit(X_train_scaled, y_train)

    # Evaluate the model on the test set
    y_test_pred = classifier.predict(X_test_scaled)

    # Calculate ROC AUC score for the classifier
    y_test_proba = classifier.predict_proba(X_test_scaled)[:, 1]
    roc_auc = roc_auc_score(y_test, y_test_proba)

    print(f"\nEvaluation for {classifier_name} completed.")
    print(f"ROC AUC Score: {roc_auc}\n")

    # Plot ROC curve
    fpr, tpr, thresholds = roc_curve(y_test, y_test_proba)
    plt.figure()
    plt.plot(fpr, tpr, color='blue', lw=2, label='ROC curve (area = %0.2f)' % roc_auc)
    plt.plot([0, 1], [0, 1], color='red', lw=2, linestyle='--')
    plt.xlim([0.0, 1.0])
    plt.ylim([0.0, 1.05])
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title(f'Receiver Operating Characteristic ({classifier_name})')
    roc_curve_plot_path = f'{base_folder}/roc_curve_{classifier_name}_seed_{seed}.png'
    plt.savefig(roc_curve_plot_path)
    plt.close()  # Close the figure

    # Confusion Matrix
    cm = confusion_matrix(y_test, y_test_pred)
    disp = ConfusionMatrixDisplay(confusion_matrix=cm, display_labels=classifier.classes_)
    confusion_matrix_plot_path = f'{base_folder}/confusion_matrix_{classifier_name}_seed_{seed}.png'
    disp.plot(cmap=plt.cm.Blues)
    plt.title(f'Confusion Matrix ({classifier_name})')
    plt.savefig(confusion_matrix_plot_path)
    plt.close()  # Close the figure

    # Create a Word document for the current classifier
    doc = Document()
    doc.add_heading(f'Diabetes Classification Results using {classifier_name}', level=1)

    # Add classification report
    doc.add_heading('Classification Report:', level=2)
    doc.add_paragraph(classification_report(y_test, y_test_pred))

    # Add ROC AUC score
    doc.add_heading('ROC AUC Score:', level=2)
    doc.add_paragraph(f'ROC AUC Score ({classifier_name}): {roc_auc}')

    # Add ROC curve plot to document
    doc.add_heading('ROC Curve:', level=2)
    doc.add_picture(roc_curve_plot_path, width=Inches(5.0))

    # Add confusion matrix
    doc.add_heading('Confusion Matrix:', level=2)
    doc.add_paragraph(f'True Negatives: {cm[0][0]}')
    doc.add_paragraph(f'False Positives: {cm[0][1]}')
    doc.add_paragraph(f'False Negatives: {cm[1][0]}')
    doc.add_paragraph(f'True Positives: {cm[1][1]}')

    # Add confusion matrix plot to document
    doc.add_heading('Confusion Matrix Plot:', level=2)
    doc.add_picture(confusion_matrix_plot_path, width=Inches(5.0))

    # Save the document for the current classifier
    classifier_doc_path = f'{base_folder}/{classifier_name}_seed_{seed}_supervised.docx'
    doc.save(classifier_doc_path)

    # Append the results to the overall document
    overall_doc.add_heading(f'Results for {classifier_name} (Seed {seed})', level=2)
    overall_doc.add_paragraph(f'ROC AUC Score: {roc_auc}')
    overall_doc.add_paragraph(f'Document saved as: {classifier_doc_path}')

# Loop through all seeds
for seed in seeds:
    # Loop through all classifiers
    for classifier_name, classifier in classifiers.items():
        base_folder = f'{classifier_name}_Supervised'  # Base folder for each classifier
        os.makedirs(base_folder, exist_ok=True)  # Create the base folder
        run_supervised_learning_process(classifier_name, classifier, seed, base_folder)

# Save the consolidated overall document
overall_doc_path = 'diabetes_classification_overall_results_supervised.docx'
overall_doc.save(overall_doc_path)

print("All classifier results have been saved successfully!")
print(f"Overall document saved as: {overall_doc_path}")
